import tkinter as tk
from tkinter import filedialog, simpledialog
import os
from datetime import datetime
from docx import Document
import pypandoc

def join_text_files():
    # Create the root window
    root = tk.Tk()
    root.withdraw()  # Hide the root window

    # Open file dialog to select multiple files
    file_paths = filedialog.askopenfilenames(title="Select text, markdown, HTML, DOCX, RTF, or RTFD files to join",
                                             filetypes=[("Text, Markdown, HTML, DOCX, RTF, and RTFD Files", "*.txt *.md *.html *.docx *.rtf *.rtfd")])

    if not file_paths:
        print("No files selected. Exiting.")
        return

    # Ask user for the output format
    output_format = simpledialog.askstring("Output Format", "Enter output format (txt, md, html, docx, rtf, or rtfd):")
    if output_format not in ["txt", "md", "html", "docx", "rtf", "rtfd"]:
        print("Invalid format selected. Exiting.")
        return

    # Prepare the output file name
    file_count = len(file_paths)
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    output_file_name = f"Merged_text_of_{file_count}_files_{timestamp}.{output_format}"
    output_dir = os.path.dirname(file_paths[0])
    output_file_path = os.path.join(output_dir, output_file_name)

    # Join the files
    if output_format == "docx":
        output_doc = Document()
        for file_path in file_paths:
            file_name = os.path.basename(file_path)
            output_doc.add_heading(file_name, level=1)
            if file_path.endswith(".docx"):
                input_doc = Document(file_path)
                for para in input_doc.paragraphs:
                    output_doc.add_paragraph(para.text)
            else:
                content = pypandoc.convert_file(file_path, 'plain')
                output_doc.add_paragraph(content)
            output_doc.add_paragraph("\n")  # Add one free line after text
        output_doc.save(output_file_path)
    else:
        content_list = []
        for file_path in file_paths:
            file_name = os.path.basename(file_path)
            content = pypandoc.convert_file(file_path, 'plain')
            content_list.append(f"\n\n{file_name}\n{content}\n")

        merged_content = "\n".join(content_list)
        pypandoc.convert_text(merged_content, output_format, format='md', outputfile=output_file_path)

    print(f"Files have been successfully merged into {output_file_path}")

if __name__ == "__main__":
    join_text_files()